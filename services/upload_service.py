"""
Upload Service — Parse PDF/DOCX, chunk, embed, and store in isolated ChromaDB collection.

Features:
- PDF text extraction via pdfplumber
- DOCX text extraction via python-docx (docx.Document)
- Chunking via LangChain RecursiveCharacterTextSplitter
- Embedding via remote API (same as ingestion_service)
- ChromaDB access via SUBPROCESS ISOLATION to prevent FFI segfaults on Windows
- Per-file progress tracking via in-memory store
- Full traceback logging on every exception
"""

import os
import sys
import uuid
import json
import asyncio
import logging
import hashlib
import tempfile
import traceback
import subprocess
from pathlib import Path

import httpx
import pdfplumber

# ── python-docx import (NOT the legacy 'docx' package) ──
try:
    from docx import Document as DocxDocument
except ImportError as _ie:
    logging.getLogger(__name__).error(
        "CRITICAL: 'python-docx' is not installed or importable. "
        "Run: pip install python-docx\n" + traceback.format_exc()
    )
    raise

# ── lxml check (python-docx depends on it) ──
try:
    import lxml  # noqa: F401
except ImportError:
    logging.getLogger(__name__).warning(
        "lxml is not installed — python-docx may fail. "
        "Run: pip install lxml"
    )

from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

# Path to the subprocess helper script
_HELPER_SCRIPT = str(
    Path(__file__).parent / "chroma_upload_helper.py"
)

# ─── In-memory progress store ───────────────────────────────────
_progress_store: dict[str, dict] = {}

# ─── Subprocess lock (serialises all ChromaDB subprocess calls) ─
_chroma_write_lock = asyncio.Lock()


# ─── Subprocess ChromaDB caller ─────────────────────────────────

def _call_chroma_subprocess(request_data: dict, timeout: int = 60) -> dict:
    """
    Run a ChromaDB operation in a SEPARATE PROCESS to avoid
    C-level FFI/Rust segfaults on Windows.

    Sends JSON via stdin, reads JSON from stdout.
    """
    try:
        input_bytes = json.dumps(request_data, ensure_ascii=False).encode('utf-8')
        env = {**os.environ, 'PYTHONIOENCODING': 'utf-8'}
        result = subprocess.run(
            [sys.executable, _HELPER_SCRIPT],
            input=input_bytes,
            capture_output=True,
            timeout=timeout,
            env=env
        )

        if result.returncode != 0:
            stderr_text = result.stderr.decode('utf-8', errors='replace') if result.stderr else ''
            logger.error(
                f"[ChromaDB Sub] Process exited with code {result.returncode}\n"
                f"stderr: {stderr_text}"
            )
            return {"error": f"Subprocess failed: {stderr_text[:500]}"}

        stdout_text = result.stdout.decode('utf-8', errors='replace') if result.stdout else ''
        if not stdout_text.strip():
            logger.error("[ChromaDB Sub] Empty stdout from subprocess")
            return {"error": "Empty response from subprocess"}

        return json.loads(stdout_text.strip())

    except subprocess.TimeoutExpired:
        logger.error(f"[ChromaDB Sub] Timeout ({timeout}s) for action={request_data.get('action')}")
        return {"error": "Subprocess timeout"}
    except Exception as e:
        logger.error(f"[ChromaDB Sub] Error: {e}\n{traceback.format_exc()}")
        return {"error": str(e)}


# ─── Document Parsing ───────────────────────────────────────────

def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            logger.info(f"[PDF] Opened '{file_path}', pages={len(pdf.pages)}")
            for idx, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    logger.debug(f"[PDF] Page {idx} returned no text")
    except Exception as e:
        logger.error(f"[PDF] FAILED to parse '{file_path}': {e}\n{traceback.format_exc()}")
        raise
    return "\n\n".join(text_parts)


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx (docx.Document)."""
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        logger.info(f"[DOCX] Opened '{file_path}', paragraphs={len(paragraphs)}, "
                     f"tables={len(doc.tables)}")

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
    except Exception as e:
        logger.error(f"[DOCX] FAILED to parse '{file_path}': {e}\n{traceback.format_exc()}")
        raise
    return "\n\n".join(paragraphs)


def parse_document(file_path: str, filename: str) -> str:
    """Parse a document based on file extension."""
    ext = Path(filename).suffix.lower()
    logger.info(f"[Parse] file='{filename}', ext='{ext}', path='{file_path}'")
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ─── Chunking ───────────────────────────────────────────────────

def chunk_document(text: str, chunk_size: int = 1000,
                   chunk_overlap: int = 200) -> list[str]:
    """Split document text into chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    return splitter.split_text(text)


# ─── Embedding ──────────────────────────────────────────────────

async def _get_embeddings(texts: list[str], api_url: str, model: str,
                          timeout: int = 120) -> list[list[float]]:
    """Call remote embedding API (Ollama /api/embed format)."""
    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            resp = await client.post(api_url, json={
                "model": model,
                "input": texts
            })
            resp.raise_for_status()
            data = resp.json()
            embeddings = data.get("embeddings", [])
            if not embeddings:
                logger.error(f"[Embed] API returned empty embeddings. response={data}")
            return embeddings
    except Exception as e:
        logger.error(f"[Embed] FAILED: url={api_url}, model={model}, "
                     f"batch_size={len(texts)}\n{traceback.format_exc()}")
        raise


# ─── ChromaDB Operations (via subprocess) ───────────────────────

def get_uploaded_file_count(config: dict) -> int:
    """Return the number of unique source files (via subprocess)."""
    upload_cfg = config["upload_rag_settings"]
    chromadb_path = str(Path(config["local_storage"]["chromadb_path"]).resolve())
    collection_name = upload_cfg["collection_name"]

    result = _call_chroma_subprocess({
        "action": "count_files",
        "chromadb_path": chromadb_path,
        "collection_name": collection_name
    })
    return result.get("count", 0)


def get_uploaded_files(config: dict) -> list[str]:
    """Return list of unique source filenames (via subprocess)."""
    upload_cfg = config["upload_rag_settings"]
    chromadb_path = str(Path(config["local_storage"]["chromadb_path"]).resolve())
    collection_name = upload_cfg["collection_name"]

    result = _call_chroma_subprocess({
        "action": "list_files",
        "chromadb_path": chromadb_path,
        "collection_name": collection_name
    })
    return result.get("files", [])


def reset_collection(config: dict) -> None:
    """Delete and recreate the upload collection (via subprocess)."""
    upload_cfg = config["upload_rag_settings"]
    chromadb_path = str(Path(config["local_storage"]["chromadb_path"]).resolve())
    collection_name = upload_cfg["collection_name"]

    result = _call_chroma_subprocess({
        "action": "reset",
        "chromadb_path": chromadb_path,
        "collection_name": collection_name
    })
    if result.get("error"):
        raise RuntimeError(f"Reset failed: {result['error']}")
    logger.info(f"Collection '{collection_name}' reset via subprocess")


def delete_doc(config: dict, filename: str) -> None:
    """Delete all chunks for a specific source file (via subprocess)."""
    upload_cfg = config["upload_rag_settings"]
    chromadb_path = str(Path(config["local_storage"]["chromadb_path"]).resolve())
    collection_name = upload_cfg["collection_name"]

    result = _call_chroma_subprocess({
        "action": "delete_doc",
        "chromadb_path": chromadb_path,
        "collection_name": collection_name,
        "filename": filename
    })
    if not result.get("ok"):
        raise RuntimeError(f"Delete failed: {result.get('error', 'Unknown error')}")
    logger.info(f"Deleted doc '{filename}' from '{collection_name}' via subprocess")


def _upsert_via_subprocess(chromadb_path: str, collection_name: str,
                           ids: list, documents: list,
                           embeddings: list, metadatas: list):
    """Upsert chunks to ChromaDB via subprocess."""
    result = _call_chroma_subprocess({
        "action": "upsert",
        "chromadb_path": chromadb_path,
        "collection_name": collection_name,
        "ids": ids,
        "documents": documents,
        "embeddings": embeddings,
        "metadatas": metadatas
    }, timeout=120)

    if not result.get("ok"):
        error = result.get("error", "Unknown upsert error")
        logger.error(f"[ChromaDB] Subprocess upsert failed: {error}")
        raise RuntimeError(f"ChromaDB upsert failed: {error}")
    logger.info(f"[ChromaDB] Upserted {result.get('count', 0)} chunks via subprocess")


# ─── Progress Management ────────────────────────────────────────

def create_task(filenames: list[str]) -> str:
    task_id = str(uuid.uuid4())[:8]
    _progress_store[task_id] = {
        "files": {
            fn: {"progress": 0, "status": "pending", "error": ""}
            for fn in filenames
        }
    }
    return task_id


def update_progress(task_id: str, filename: str, progress: int,
                    status: str = "processing", error: str = ""):
    if task_id in _progress_store:
        _progress_store[task_id]["files"][filename] = {
            "progress": min(progress, 100),
            "status": status,
            "error": error
        }


def get_progress(task_id: str) -> dict | None:
    return _progress_store.get(task_id)


def cleanup_task(task_id: str):
    _progress_store.pop(task_id, None)


# ─── Main Processing Pipeline ──────────────────────────────────

async def process_upload(file_path: str, filename: str, config: dict,
                         task_id: str) -> None:
    """
    Full upload pipeline for a single file:
    1. Parse document → text
    2. Chunk text
    3. Embed chunks (in batches)
    4. Upsert to ChromaDB via subprocess
    """
    upload_cfg = config["upload_rag_settings"]
    embedding_api_url = config["embedding_api_url"]
    embedding_model = config["rag_parameters"]["embedding_model"]
    chromadb_path = str(Path(config["local_storage"]["chromadb_path"]).resolve())
    collection_name = upload_cfg["collection_name"]
    chunk_size = upload_cfg.get("chunk_size", 1000)
    chunk_overlap = upload_cfg.get("chunk_overlap", 200)

    # ── Debug: log file metadata before processing ──
    try:
        file_size = os.path.getsize(file_path)
        file_readable = os.access(file_path, os.R_OK)
        dir_writable = os.access(os.path.dirname(file_path), os.W_OK)
        logger.info(
            f"[Upload] START '{filename}' | "
            f"path={file_path} | size={file_size} bytes | "
            f"readable={file_readable} | dir_writable={dir_writable}"
        )
    except Exception as e:
        logger.warning(f"[Upload] Could not stat file '{file_path}': {e}")

    try:
        # Step 1: Parse (0% → 10%)
        update_progress(task_id, filename, 5, "processing")
        text = parse_document(file_path, filename)
        if not text.strip():
            update_progress(task_id, filename, 100, "error",
                            "No text content found in document")
            return
        update_progress(task_id, filename, 10, "processing")
        logger.info(f"[Upload] Parsed '{filename}': {len(text)} chars")

        # Step 2: Chunk (10% → 20%)
        chunks = chunk_document(text, chunk_size, chunk_overlap)
        if not chunks:
            update_progress(task_id, filename, 100, "error",
                            "No chunks generated from document")
            return
        update_progress(task_id, filename, 20, "processing")
        logger.info(f"[Upload] Chunked '{filename}': {len(chunks)} chunks")

        # Step 3: Embed in batches (20% → 90%)
        batch_size = 10
        all_embeddings = []
        total_batches = (len(chunks) + batch_size - 1) // batch_size

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            batch_idx = i // batch_size
            logger.debug(f"[Upload] Embedding batch {batch_idx+1}/{total_batches} "
                         f"for '{filename}'")
            embeddings = await _get_embeddings(
                batch, embedding_api_url, embedding_model
            )
            all_embeddings.extend(embeddings)

            pct = 20 + int(70 * (batch_idx + 1) / total_batches)
            update_progress(task_id, filename, pct, "processing")

        logger.info(f"[Upload] Embedded '{filename}': {len(all_embeddings)} vectors")

        # Step 4: Upsert to ChromaDB via subprocess (90% → 100%)
        update_progress(task_id, filename, 90, "processing")

        file_hash = hashlib.md5(filename.encode()).hexdigest()[:8]
        ids = [f"upload_{file_hash}_{j}" for j in range(len(chunks))]
        metadatas = [{"source_file": filename, "chunk_index": j}
                     for j in range(len(chunks))]

        async with _chroma_write_lock:
            # Run upsert in subprocess — ChromaDB Rust FFI crashes
            # in the main process on Windows
            _upsert_via_subprocess(chromadb_path, collection_name,
                                   ids, chunks, all_embeddings, metadatas)

        update_progress(task_id, filename, 100, "done")
        logger.info(f"[Upload] COMPLETED '{filename}' — "
                     f"{len(chunks)} chunks stored in '{collection_name}'")

    except Exception as e:
        logger.error(
            f"[Upload] FAILED processing '{filename}': {e}\n"
            f"{traceback.format_exc()}"
        )
        update_progress(task_id, filename, 0, "error", str(e))
    finally:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.debug(f"[Upload] Removed temp file '{file_path}'")
        except Exception as e:
            logger.warning(f"[Upload] Could not remove temp file '{file_path}': {e}")
